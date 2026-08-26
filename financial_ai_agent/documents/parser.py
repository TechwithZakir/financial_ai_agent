from __future__ import annotations

import csv
from io import StringIO
from pathlib import Path

MAX_FILE_SIZE = 25 * 1024 * 1024
MAX_CELLS = 100_000


def parse_file(path: str, original_filename: str | None = None) -> dict:
    file_path = Path(path)
    if not file_path.is_file():
        raise ValueError("Document file was not found")
    if file_path.stat().st_size > MAX_FILE_SIZE:
        raise ValueError("Document exceeds the 25 MB processing limit")
    extension = Path(original_filename or path).suffix.lower()
    parsers = {
        ".pdf": _pdf,
        ".docx": _docx,
        ".xlsx": _xlsx,
        ".csv": _csv,
        ".txt": _text,
        ".md": _text,
    }
    if extension in {".png", ".jpg", ".jpeg"}:
        return {"text": "", "pages": [], "requires_vision": True, "extension": extension}
    if extension not in parsers:
        raise ValueError(f"Unsupported document type: {extension or 'unknown'}")
    result = parsers[extension](file_path)
    result.update({"requires_vision": False, "extension": extension})
    return result


def _pdf(path: Path) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError("Password-protected PDFs are not supported")
    pages = [{"page": index + 1, "text": page.extract_text() or ""}
             for index, page in enumerate(reader.pages)]
    return {"text": "\n\n".join(page["text"] for page in pages), "pages": pages}


def _docx(path: Path) -> dict:
    from docx import Document
    document = Document(str(path))
    lines = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        lines.extend(" | ".join(cell.text.strip() for cell in row.cells) for row in table.rows)
    text = "\n".join(lines)
    return {"text": text, "pages": [{"page": None, "text": text}]}


def _xlsx(path: Path) -> dict:
    from openpyxl import load_workbook
    workbook = load_workbook(str(path), read_only=True, data_only=True)
    lines, cells = [], 0
    for sheet in workbook.worksheets:
        lines.append(f"Sheet: {sheet.title}")
        for row in sheet.iter_rows(values_only=True):
            cells += len(row)
            if cells > MAX_CELLS:
                raise ValueError("Spreadsheet exceeds the 100,000-cell processing limit")
            values = ["" if value is None else str(value) for value in row]
            if any(values):
                lines.append(" | ".join(values))
    text = "\n".join(lines)
    return {"text": text, "pages": [{"page": None, "text": text}]}


def _csv(path: Path) -> dict:
    content = path.read_text(encoding="utf-8-sig", errors="replace")
    rows = list(csv.reader(StringIO(content)))
    if sum(len(row) for row in rows) > MAX_CELLS:
        raise ValueError("CSV exceeds the 100,000-cell processing limit")
    text = "\n".join(" | ".join(row) for row in rows)
    return {"text": text, "pages": [{"page": None, "text": text}]}


def _text(path: Path) -> dict:
    text = path.read_text(encoding="utf-8-sig", errors="replace")
    return {"text": text, "pages": [{"page": None, "text": text}]}

