from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_RIGHT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

OUTPUT = Path(__file__).resolve().parents[1] / "sample_data"
OUTPUT.mkdir(parents=True, exist_ok=True)

COMPANY = "Northstar Trading Ltd. (Synthetic Demo)"
PERIOD = "Year ended 31 December 2025"


def build_pdf():
    path = OUTPUT / "synthetic_financial_statements_2025.pdf"
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Right", parent=styles["BodyText"], alignment=TA_RIGHT))
    doc = SimpleDocTemplate(str(path), pagesize=A4, rightMargin=18*mm, leftMargin=18*mm,
                            topMargin=18*mm, bottomMargin=18*mm, title=COMPANY)
    story = [Paragraph(COMPANY, styles["Title"]), Paragraph(PERIOD, styles["Heading2"]),
             Paragraph("SYNTHETIC TEST DATA - NOT A REAL COMPANY", styles["Heading3"]), Spacer(1, 7*mm)]
    for title, rows in statements():
        story.append(Paragraph(title, styles["Heading2"]))
        data = [["Line item", "2025 (BDT)", "2024 (BDT)"]] + rows
        table = Table(data, colWidths=[92*mm, 40*mm, 40*mm], repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), colors.HexColor("#17365D")),
            ("TEXTCOLOR", (0,0), (-1,0), colors.white), ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
            ("ALIGN", (1,1), (-1,-1), "RIGHT"), ("LINEBELOW", (0,0), (-1,0), 1, colors.HexColor("#17365D")),
            ("LINEABOVE", (0,-1), (-1,-1), 1, colors.HexColor("#17365D")),
            ("FONTNAME", (0,-1), (-1,-1), "Helvetica-Bold"),
            ("ROWBACKGROUNDS", (0,1), (-1,-2), [colors.white, colors.HexColor("#F3F6FA")]),
            ("BOTTOMPADDING", (0,0), (-1,-1), 6), ("TOPPADDING", (0,0), (-1,-1), 6),
        ]))
        story.extend([table, Spacer(1, 8*mm)])
    story.append(Paragraph("All names, amounts, identifiers, and transactions are fictional and intended only for software testing.", styles["BodyText"]))
    doc.build(story)


def build_docx():
    path = OUTPUT / "synthetic_credit_application_summary.docx"
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = Inches(0.75)
    section.left_margin = section.right_margin = Inches(0.8)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Synthetic Credit Application Summary")
    run.bold = True; run.font.size = Pt(22); run.font.color.rgb = RGBColor(23, 54, 93)
    subtitle = doc.add_paragraph("Northstar Trading Ltd. - Test Case FA-2025-001")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("SYNTHETIC TEST DATA - NOT A REAL APPLICATION").runs[0].bold = True
    doc.add_heading("Applicant overview", level=1)
    for label, value in [("Applicant", "Northstar Trading Ltd."), ("Requested facility", "BDT 8,000,000 working-capital line"),
                         ("Term", "24 months"), ("Purpose", "Inventory purchases and receivables financing")]:
        p = doc.add_paragraph(); p.add_run(f"{label}: ").bold = True; p.add_run(value)
    doc.add_heading("Financial snapshot", level=1)
    table = doc.add_table(rows=1, cols=3); table.style = "Light Shading Accent 1"
    for cell, text in zip(table.rows[0].cells, ["Metric", "2025", "2024"]): cell.text = text
    for row in [("Revenue", "BDT 48,500,000", "BDT 41,200,000"), ("EBITDA", "BDT 6,900,000", "BDT 5,400,000"),
                ("Net income", "BDT 3,650,000", "BDT 2,710,000"), ("Total debt", "BDT 9,500,000", "BDT 8,100,000")]:
        cells = table.add_row().cells
        for cell, text in zip(cells, row): cell.text = text
    doc.add_heading("Documents supplied", level=1)
    for item in ["Audited financial statements for 2024 and 2025", "Six-month bank statement", "Trade license", "Tax return acknowledgement"]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_heading("Review flags", level=1)
    for item in ["Customer concentration exceeds 30% for one buyer", "Receivable days increased from 48 to 61", "Latest inventory ageing report is missing"]:
        doc.add_paragraph(item, style="List Bullet")
    footer = section.footer.paragraphs[0]; footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("Synthetic fixture for Financial AI Agent testing")
    doc.save(path)


def statements():
    return [
        ("Statement of profit or loss", [
            ["Revenue", "48,500,000", "41,200,000"], ["Cost of sales", "(32,250,000)", "(28,400,000)"],
            ["Gross profit", "16,250,000", "12,800,000"], ["Operating expenses", "(9,350,000)", "(7,400,000)"],
            ["Operating profit", "6,900,000", "5,400,000"], ["Finance costs", "(1,300,000)", "(1,050,000)"],
            ["Tax expense", "(1,950,000)", "(1,640,000)"], ["Net income", "3,650,000", "2,710,000"],
        ]),
        ("Statement of financial position", [
            ["Cash and cash equivalents", "4,200,000", "3,100,000"], ["Trade receivables", "8,600,000", "6,900,000"],
            ["Inventory", "7,400,000", "6,500,000"], ["Property and equipment", "12,300,000", "11,900,000"],
            ["Total assets", "32,500,000", "28,400,000"], ["Trade payables", "5,100,000", "4,700,000"],
            ["Borrowings", "9,500,000", "8,100,000"], ["Other liabilities", "2,250,000", "2,100,000"],
            ["Total liabilities", "16,850,000", "14,900,000"], ["Total equity", "15,650,000", "13,500,000"],
        ]),
        ("Statement of cash flows", [
            ["Net cash from operating activities", "5,250,000", "4,100,000"],
            ["Net cash used in investing activities", "(2,100,000)", "(1,850,000)"],
            ["Net cash used in financing activities", "(2,050,000)", "(1,600,000)"],
            ["Net increase in cash", "1,100,000", "650,000"],
        ]),
    ]


if __name__ == "__main__":
    build_pdf(); build_docx()

